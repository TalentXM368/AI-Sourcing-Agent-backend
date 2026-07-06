#!/usr/bin/env python3
"""Bulk resume and job description parser.

What it does:
- Extracts text from resumes/JDs in bulk
- Parses locally with rules first
- Uses OpenAI fallback only when local parsing is weak and an API key is available
- Writes one JSON output file with parsed results

Run:
    python main.py                          # Parse resumes (default)
    python main.py --parse-jds              # Parse job descriptions
    python main.py --parse-jds --input-dir ./jd --output-file ./parsed_jds.json

Default workspace routes:
- resumes folder: ./resumes
- jd folder: ./jd
- resume output: ./parsed_resumes.json
- jd output: ./parsed_jds.json
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import logging
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from bs4 import BeautifulSoup
from docx import Document
from odf import teletype
from odf.opendocument import load as odf_load
from pypdf import PdfReader
from striprtf.striprtf import rtf_to_text

# Ensure project root is on sys.path so imports work when running from backend/
PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None

try:
    from backend.parsers import parse_jd as parse_jd_file, process_directory as process_jd_directory, extract_text_from_file as extract_jd_text
    from backend.parsers import JDLLMConfig, OpenAIJDParser, ClientContextParser
except Exception:  # pragma: no cover
    try:
        from parsers import parse_jd as parse_jd_file, process_directory as process_jd_directory, extract_text_from_file as extract_jd_text
        from parsers import JDLLMConfig, OpenAIJDParser, ClientContextParser
    except Exception:
        parse_jd_file = None
        process_jd_directory = None
        extract_jd_text = None
        JDLLMConfig = None
        OpenAIJDParser = None
        ClientContextParser = None

try:
    from backend.pipeline import VectorPipeline
except Exception:  # pragma: no cover
    try:
        from pipeline import VectorPipeline
    except Exception:
        VectorPipeline = None


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".rtf",
    ".html",
    ".htm",
    ".odt",
    ".json",
    ".csv",
}

DEFAULT_OUTPUT_SCHEMA = {
    "full_name": None,
    "email": None,
    "phone": None,
    "linkedin_url": None,
    "github_url": None,
    "skills": [],
    "projects": [],
    "total_experience_years": None,
    "education": None,
    "current_role": None,
}

SKILL_NORMALIZATION = {
    "ml": "machine learning",
    "ai": "artificial intelligence",
    "nlp": "natural language processing",
    "dl": "deep learning",
    "genai": "generative ai",
    "llm": "large language models",
    "sql": "sql",
    "aws": "amazon web services",
    "gcp": "google cloud platform",
    "azure": "microsoft azure",
}

KNOWN_JOB_TITLES = {
    "software engineer",
    "senior software engineer",
    "full stack developer",
    "backend developer",
    "frontend developer",
    "data scientist",
    "data analyst",
    "machine learning engineer",
    "devops engineer",
    "cloud engineer",
    "project manager",
    "product manager",
    "business analyst",
    "qa engineer",
    "test engineer",
}

EDUCATION_PATTERNS = [
    (4, r"\b(ph\.?d|doctorate)\b"),
    (3, r"\b(m\.?tech|m\.?s\.?|master(?:'s)?|mba)\b"),
    (2, r"\b(b\.?tech|b\.?e\.?|b\.?s\.?|bachelor(?:'s)?)\b"),
    (1, r"\b(diploma|associate degree)\b"),
]

PROJECT_SECTION_HEADERS = {
    "projects",
    "project",
    "academic projects",
    "personal projects",
    "key projects",
    "project experience",
    "selected projects",
}

URL_CLEANUP_RE = re.compile(r"[\)\],.;:\'\"<>]+$")
LINKEDIN_HOSTS = ("linkedin.com", "www.linkedin.com", "in.linkedin.com")
GITHUB_HOSTS = ("github.com", "www.github.com")

WORKSPACE_ROOT = PROJECT_ROOT
DEFAULT_INPUT_DIR = WORKSPACE_ROOT / "backend" / "resumes"
DEFAULT_OUTPUT_FILE = WORKSPACE_ROOT / "backend" / "results" / "parsed_resumes.json"
DEFAULT_CLIENT_DATA_DIR = WORKSPACE_ROOT / "backend" / "client_data"
DEFAULT_PARSED_CLIENT_DATA_FILE = WORKSPACE_ROOT / "backend" / "results" / "parsed_client_data.json"


def load_env_file(env_path: Path) -> None:
    if not env_path.exists() or not env_path.is_file():
        return
    for raw_line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def load_client_context(client_data_dir: Path) -> Optional[Dict[str, Any]]:
    """Load and merge client context files from the flat client_data folder.

    Returns a bundle containing the parsed source records and the merged
    context used for matching.
    """
    if not client_data_dir.exists() or not client_data_dir.is_dir():
        return None

    context_files = sorted(
        path for path in client_data_dir.iterdir()
        if path.is_file()
        and path.name.lower() != "readme.md"
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    if not context_files:
        return None

    parsed_sources: List[Dict[str, Any]] = []
    contexts: List[Dict[str, Any]] = []
    parser = ClientContextParser() if ClientContextParser else None

    for context_file in context_files:
        source_record: Dict[str, Any] = {
            "file_name": context_file.name,
            "file_path": str(context_file),
            "status": "skipped",
            "parsed_context": None,
        }

        raw_context: Optional[Dict[str, Any]] = None
        try:
            raw_loaded = json.loads(context_file.read_text(encoding="utf-8"))
            if isinstance(raw_loaded, dict):
                raw_context = raw_loaded
            else:
                raw_context = {"Description": str(raw_loaded)}
        except Exception:
            try:
                raw_text = ResumeTextExtractor.extract_text(context_file)
            except Exception as exc:
                logging.warning("Skipping client data file %s: %s", context_file.name, exc)
                source_record["error"] = str(exc)
                parsed_sources.append(source_record)
                continue

            raw_context = {
                "Account_Name": context_file.stem,
                "Description": raw_text,
                "Source_File": context_file.name,
                "Source_Format": context_file.suffix.lower().lstrip("."),
            }

        if isinstance(raw_context, dict) and {
            "client_name",
            "culture",
            "hiring_preferences",
            "role_context",
        }.issubset(raw_context.keys()):
            context = dict(raw_context)
        elif parser and isinstance(raw_context, dict):
            context = parser.parse_zoho_profile(raw_context)
        else:
            logging.warning("Skipping unsupported client data file %s", context_file.name)
            source_record["error"] = "Unsupported client data format"
            parsed_sources.append(source_record)
            continue

        context.setdefault("_source_files", [])
        context["_source_files"].append(context_file.name)
        contexts.append(context)
        source_record["status"] = "parsed"
        source_record["parsed_context"] = context
        parsed_sources.append(source_record)

    if not contexts:
        return None

    if len(contexts) == 1 or parser is None:
        merged = contexts[0]
    else:
        merged = parser.merge_contexts(contexts)

    merged["_source_files"] = [path.name for path in context_files]
    quality_summary = {
        "source_file_count": len(context_files),
        "parsed_source_count": sum(1 for item in parsed_sources if item.get("status") == "parsed"),
        "skipped_source_count": sum(1 for item in parsed_sources if item.get("status") != "parsed"),
        "has_common_skills": bool(merged.get("historical_patterns", {}).get("common_skills")),
        "has_industry": bool(merged.get("role_context", {}).get("industry")) and merged.get("role_context", {}).get("industry") != "unknown",
        "has_preferences": bool(merged.get("hiring_preferences", {}).get("preferred_background")),
    }
    logging.info(
        "Loaded client context bundle: %d sources, %d parsed, %d skipped",
        quality_summary["source_file_count"],
        quality_summary["parsed_source_count"],
        quality_summary["skipped_source_count"],
    )
    return {
        "merged_context": merged,
        "parsed_sources": parsed_sources,
        "source_files": [path.name for path in context_files],
        "quality_summary": quality_summary,
    }


class ResumeTextExtractor:
    @staticmethod
    def extract_text(file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return ResumeTextExtractor._from_pdf(file_path)
        if suffix == ".docx":
            return ResumeTextExtractor._from_docx(file_path)
        if suffix in {".txt", ".md", ".json", ".csv"}:
            return ResumeTextExtractor._from_text(file_path)
        if suffix == ".rtf":
            return ResumeTextExtractor._from_rtf(file_path)
        if suffix in {".html", ".htm"}:
            return ResumeTextExtractor._from_html(file_path)
        if suffix == ".odt":
            return ResumeTextExtractor._from_odt(file_path)
        return ResumeTextExtractor._from_text(file_path)

    @staticmethod
    def _from_pdf(file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        link_texts = []
        for page in reader.pages:
            link_texts.extend(ResumeTextExtractor._extract_pdf_links(page))
        pages.extend(link_texts)
        return "\n".join(pages).strip()

    @staticmethod
    def _from_docx(file_path: Path) -> str:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs]
        hyperlinks = ResumeTextExtractor._extract_docx_hyperlinks(doc)
        paragraphs.extend(hyperlinks)
        return "\n".join(paragraphs).strip()

    @staticmethod
    def _from_text(file_path: Path) -> str:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return file_path.read_text(encoding=encoding).strip()
            except UnicodeDecodeError:
                continue
        return file_path.read_text(errors="ignore").strip()

    @staticmethod
    def _from_rtf(file_path: Path) -> str:
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        return rtf_to_text(raw).strip()

    @staticmethod
    def _from_html(file_path: Path) -> str:
        html = file_path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        texts = [soup.get_text(separator="\n", strip=True)]
        texts.extend(a.get("href", "") for a in soup.find_all("a", href=True))
        return "\n".join(texts).strip()

    @staticmethod
    def _from_odt(file_path: Path) -> str:
        doc = odf_load(str(file_path))
        return teletype.extractText(doc).strip()

    @staticmethod
    def _extract_docx_hyperlinks(doc: Document) -> List[str]:
        hyperlinks: List[str] = []
        rels = doc.part.rels

        for paragraph in doc.paragraphs:
            para = paragraph._p
            for hyperlink in para.xpath('.//w:hyperlink'):
                rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if not rel_id or rel_id not in rels:
                    continue
                target = rels[rel_id].target_ref
                if target:
                    hyperlinks.append(target)

        return hyperlinks

    @staticmethod
    def _extract_pdf_links(page: Any) -> List[str]:
        links: List[str] = []
        annotations = page.get('/Annots') or []
        for annotation_ref in annotations:
            try:
                annotation = annotation_ref.get_object()
                action = annotation.get('/A')
                if action and action.get('/URI'):
                    links.append(str(action.get('/URI')))
            except Exception:
                continue
        return links


class RuleBasedResumeParser:
    def parse_resume(self, resume_text: str) -> Dict[str, Any]:
        return {
            "full_name": self._extract_full_name(resume_text),
            "email": self._extract_email(resume_text),
            "phone": self._extract_phone(resume_text),
            "linkedin_url": self._extract_linkedin_url(resume_text),
            "github_url": self._extract_github_url(resume_text),
            "skills": self._extract_skills(resume_text),
            "projects": self._extract_projects(resume_text),
            "total_experience_years": self._extract_experience_years(resume_text),
            "education": self._extract_highest_education(resume_text),
            "current_role": self._extract_current_role(resume_text),
        }

    @staticmethod
    def _extract_full_name(text: str) -> Optional[str]:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for line in lines[:20]:
            if len(line) > 60:
                continue
            if re.search(r"[@\d]|https?://|www\.", line.lower()):
                continue
            parts = re.split(r"\s+", line)
            if len(parts) < 2 or len(parts) > 4:
                continue
            if all(re.fullmatch(r"[A-Za-z][A-Za-z\-'.]*", p) for p in parts):
                return " ".join(parts)
        return None

    @staticmethod
    def _extract_email(text: str) -> Optional[str]:
        matches = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
        return matches[0].lower() if matches else None

    @staticmethod
    def _extract_phone(text: str) -> Optional[str]:
        candidates = re.findall(r"(?:\+?\d[\d\s().-]{7,}\d)", text)
        for cand in candidates:
            digits = re.sub(r"[^\d+]", "", cand)
            if len(re.sub(r"\D", "", digits)) >= 7:
                return digits
        return None

    @staticmethod
    def _extract_linkedin_url(text: str) -> Optional[str]:
        return RuleBasedResumeParser._extract_social_url(text, LINKEDIN_HOSTS)

    @staticmethod
    def _extract_github_url(text: str) -> Optional[str]:
        return RuleBasedResumeParser._extract_social_url(text, GITHUB_HOSTS)

    @staticmethod
    def _extract_social_url(text: str, allowed_hosts: tuple[str, ...]) -> Optional[str]:
        candidates = re.findall(r"https?://[^\s)\],>]+|www\.[^\s)\],>]+", text, flags=re.IGNORECASE)
        for candidate in candidates:
            cleaned = RuleBasedResumeParser._normalize_url(candidate)
            if cleaned and RuleBasedResumeParser._host_matches(cleaned, allowed_hosts):
                return cleaned

        # Fallback for hyperlinked labels like "LinkedIn" or "GitHub" where the visible text is not the URL.
        label_patterns = {
            LINKEDIN_HOSTS: r"(?:linkedin\s*[:\-]\s*|linkedin\s+profile\s*[:\-]\s*)(https?://[^\s)\],>]+|www\.[^\s)\],>]+|[A-Za-z0-9./_\-]+)",
            GITHUB_HOSTS: r"(?:github\s*[:\-]\s*|github\s+profile\s*[:\-]\s*)(https?://[^\s)\],>]+|www\.[^\s)\],>]+|[A-Za-z0-9./_\-]+)",
        }
        for hosts, pattern in label_patterns.items():
            if not set(hosts).intersection(set(allowed_hosts)):
                continue
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                cleaned = RuleBasedResumeParser._normalize_url(match.group(1))
                if cleaned and RuleBasedResumeParser._host_matches(cleaned, allowed_hosts):
                    return cleaned

        return None

    @staticmethod
    def _normalize_url(value: str) -> Optional[str]:
        cleaned = value.strip().strip('"').strip("'")
        cleaned = URL_CLEANUP_RE.sub("", cleaned)
        if cleaned.startswith("www."):
            cleaned = "https://" + cleaned
        if not cleaned.startswith(("http://", "https://")) and ("linkedin.com" in cleaned or "github.com" in cleaned):
            cleaned = "https://" + cleaned.lstrip("/")
        return cleaned or None

    @staticmethod
    def _host_matches(url: str, allowed_hosts: tuple[str, ...]) -> bool:
        lower = url.lower()
        return any(host in lower for host in allowed_hosts)

    @staticmethod
    def _extract_skills(text: str) -> List[str]:
        lowered = text.lower()
        skills = set()
        for alias, normalized in SKILL_NORMALIZATION.items():
            if re.search(rf"\b{re.escape(alias)}\b", lowered):
                skills.add(normalized)
            if re.search(rf"\b{re.escape(normalized)}\b", lowered):
                skills.add(normalized)

        common_skills = [
            "python",
            "java",
            "javascript",
            "typescript",
            "react",
            "node.js",
            "django",
            "flask",
            "fastapi",
            "spring boot",
            "docker",
            "kubernetes",
            "git",
            "mysql",
            "postgresql",
            "mongodb",
            "pandas",
            "numpy",
            "tensorflow",
            "pytorch",
        ]
        for skill in common_skills:
            if re.search(rf"\b{re.escape(skill)}\b", lowered):
                skills.add(skill)
        return sorted(skills)

    @staticmethod
    def _extract_experience_years(text: str) -> Optional[float]:
        explicit = re.search(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\b", text, flags=re.IGNORECASE)
        if explicit:
            return float(explicit.group(1))

        current_year = dt.date.today().year
        ranges = re.findall(r"\b(19\d{2}|20\d{2})\s*(?:-|to|–)\s*(present|current|19\d{2}|20\d{2})\b", text, flags=re.IGNORECASE)
        total = 0
        for start, end in ranges:
            start_y = int(start)
            end_y = current_year if end.lower() in {"present", "current"} else int(end)
            if 1900 <= start_y <= end_y <= current_year + 1:
                total += end_y - start_y
        return float(total) if total > 0 else None

    @staticmethod
    def _extract_highest_education(text: str) -> Optional[str]:
        best_rank = -1
        best_value = None
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for line in lines:
            lowered = line.lower()
            for rank, pattern in EDUCATION_PATTERNS:
                if re.search(pattern, lowered) and rank > best_rank:
                    best_rank = rank
                    best_value = line
        return best_value

    @staticmethod
    def _extract_current_role(text: str) -> Optional[str]:
        label_patterns = [
            r"current\s+role\s*[:\-]\s*(.+)",
            r"designation\s*[:\-]\s*(.+)",
            r"current\s+position\s*[:\-]\s*(.+)",
            r"title\s*[:\-]\s*(.+)",
        ]
        for pattern in label_patterns:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                candidate = match.group(1).strip()
                candidate = re.split(r"\s{2,}|\||,", candidate)[0].strip()
                if candidate:
                    return candidate

        lowered_text = text.lower()
        for title in sorted(KNOWN_JOB_TITLES, key=len, reverse=True):
            if re.search(rf"\b{re.escape(title)}\b", lowered_text):
                return title
        return None

    @staticmethod
    def _extract_projects(text: str) -> List[Dict[str, Any]]:
        lines = [ln.rstrip() for ln in text.splitlines()]
        projects: List[Dict[str, Any]] = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if RuleBasedResumeParser._is_project_heading(line):
                section_lines: List[str] = []
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if next_line and RuleBasedResumeParser._looks_like_heading(next_line):
                        break
                    if next_line:
                        section_lines.append(next_line)
                    i += 1
                projects.extend(RuleBasedResumeParser._parse_project_block(section_lines))
                continue
            i += 1
        return projects

    @staticmethod
    def _is_project_heading(line: str) -> bool:
        normalized = re.sub(r"\s+", " ", line.lower()).strip().rstrip(":")
        return normalized in PROJECT_SECTION_HEADERS

    @staticmethod
    def _looks_like_heading(line: str) -> bool:
        normalized = re.sub(r"\s+", " ", line.strip().lower()).rstrip(":")
        if len(normalized) > 80:
            return False
        if normalized in PROJECT_SECTION_HEADERS:
            return True
        heading_keywords = {
            "experience",
            "education",
            "skills",
            "certifications",
            "summary",
            "objective",
            "internship",
            "work history",
            "work experience",
            "achievements",
            "contact",
        }
        return normalized in heading_keywords or bool(re.fullmatch(r"[A-Z][A-Z\s/&-]{2,}", line.strip()))

    @staticmethod
    def _parse_project_block(section_lines: List[str]) -> List[Dict[str, Any]]:
        projects: List[Dict[str, Any]] = []
        buffer: List[str] = []

        def flush_buffer() -> None:
            if not buffer:
                return
            project_text = " ".join(buffer).strip()
            if project_text:
                parsed = RuleBasedResumeParser._parse_single_project(project_text)
                if parsed:
                    projects.append(parsed)
            buffer.clear()

        for line in section_lines:
            bullet = re.sub(r"^[•\-*\u2022]+\s*", "", line).strip()
            if not bullet:
                continue
            if RuleBasedResumeParser._looks_like_new_project_item(bullet) and buffer:
                flush_buffer()
            buffer.append(bullet)

        flush_buffer()
        return projects

    @staticmethod
    def _looks_like_new_project_item(line: str) -> bool:
        if re.match(r"^(project|app|application|website|web app|mobile app)\b", line, flags=re.IGNORECASE):
            return True
        if re.search(r"[:\-–]\s+", line):
            return True
        if re.match(r"^[A-Z][A-Za-z0-9'()\- ]{2,60}$", line) and len(line.split()) <= 10:
            return True
        return False

    @staticmethod
    def _parse_single_project(project_text: str) -> Optional[Dict[str, Any]]:
        text = re.sub(r"\s+", " ", project_text).strip()
        if not text:
            return None
        title = None
        description = text
        split_match = re.match(r"^(.{2,80}?)(?:\s*[:\-–]\s+)(.+)$", text)
        if split_match:
            title = split_match.group(1).strip()
            description = split_match.group(2).strip()
        else:
            first_sentence = re.split(r"\.(?=\s|$)", text, maxsplit=1)[0].strip()
            if len(first_sentence.split()) <= 10:
                title = first_sentence
        technologies = RuleBasedResumeParser._extract_project_technologies(text)
        if title is None:
            title = description[:80].strip()
        return {
            "title": title or None,
            "description": description or None,
            "technologies": technologies,
            "raw_text": text,
        }

    @staticmethod
    def _extract_project_technologies(text: str) -> List[str]:
        lowered = text.lower()
        tech_aliases = {
            "python": "python",
            "django": "django",
            "flask": "flask",
            "fastapi": "fastapi",
            "react": "react",
            "node": "node.js",
            "node.js": "node.js",
            "javascript": "javascript",
            "typescript": "typescript",
            "sql": "sql",
            "mysql": "mysql",
            "postgres": "postgresql",
            "postgresql": "postgresql",
            "mongodb": "mongodb",
            "docker": "docker",
            "kubernetes": "kubernetes",
            "tensorflow": "tensorflow",
            "pytorch": "pytorch",
            "nlp": "natural language processing",
            "ml": "machine learning",
            "ai": "artificial intelligence",
            "aws": "amazon web services",
            "azure": "microsoft azure",
            "gcp": "google cloud platform",
        }
        technologies = set()
        for alias, normalized in tech_aliases.items():
            if re.search(rf"\b{re.escape(alias)}\b", lowered):
                technologies.add(normalized)
        return sorted(technologies)


@dataclass
class LLMConfig:
    api_key: str
    model: str = "gpt-4.1-mini"
    api_base: Optional[str] = None
    timeout_sec: int = 120


class OpenAIResumeParser:
    def __init__(self, config: LLMConfig) -> None:
        if OpenAI is None:
            raise RuntimeError("openai package is not installed")
        self.config = config
        self.client = OpenAI(api_key=config.api_key, base_url=config.api_base, timeout=config.timeout_sec)

    def parse_resume(self, resume_text: str) -> Dict[str, Any]:
        prompt = self._build_prompt(resume_text)
        raw_json = self._extract_json_from_model(prompt)
        return self._normalize(raw_json)

    def _extract_json_from_model(self, prompt: str) -> Dict[str, Any]:
        response = self.client.chat.completions.create(
            model=self.config.model,
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You extract structured resume data. Return valid JSON only, with no code fences and no extra commentary."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content or "{}"
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as exc:
            match = re.search(r"\{[\s\S]*\}", content)
            if not match:
                raise ValueError(f"LLM did not return valid JSON: {content[:500]}") from exc
            parsed = json.loads(match.group(0))
        if not isinstance(parsed, dict):
            raise ValueError("LLM output is not a JSON object")
        return parsed

    @staticmethod
    def _build_prompt(resume_text: str) -> str:
        return f'''Extract the following information from the resume text and return ONLY JSON.

Rules:
- Do not guess missing values.
- If a field is missing, use null.
- Normalize skills to lowercase standard names.
- Remove duplicate skills.
- Include projects if they are present in the resume.

Return JSON in this exact shape:
{{
  "full_name": null,
  "email": null,
  "phone": null,
    "linkedin_url": null,
    "github_url": null,
  "skills": [],
  "projects": [],
  "total_experience_years": null,
  "education": null,
  "current_role": null
}}

Resume Text:
"""
{resume_text}
"""'''

    @staticmethod
    def _normalize(value: Dict[str, Any]) -> Dict[str, Any]:
        result = dict(DEFAULT_OUTPUT_SCHEMA)
        result["full_name"] = OpenAIResumeParser._clean_string(value.get("full_name"))
        result["email"] = OpenAIResumeParser._normalize_email(value.get("email"))
        result["phone"] = OpenAIResumeParser._normalize_phone(value.get("phone"))
        result["linkedin_url"] = OpenAIResumeParser._normalize_social_url(value.get("linkedin_url"), LINKEDIN_HOSTS)
        result["github_url"] = OpenAIResumeParser._normalize_social_url(value.get("github_url"), GITHUB_HOSTS)
        result["skills"] = OpenAIResumeParser._normalize_skills(value.get("skills"))
        result["projects"] = OpenAIResumeParser._normalize_projects(value.get("projects"))
        result["total_experience_years"] = OpenAIResumeParser._normalize_experience(value.get("total_experience_years"))
        result["education"] = OpenAIResumeParser._clean_string(value.get("education"))
        result["current_role"] = OpenAIResumeParser._clean_string(value.get("current_role"))
        return result

    @staticmethod
    def _clean_string(value: Any) -> Optional[str]:
        if value is None:
            return None
        if not isinstance(value, str):
            value = str(value)
        value = value.strip()
        return value or None

    @staticmethod
    def _normalize_email(value: Any) -> Optional[str]:
        text = OpenAIResumeParser._clean_string(value)
        if not text:
            return None
        match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
        return match.group(0).lower() if match else None

    @staticmethod
    def _normalize_phone(value: Any) -> Optional[str]:
        text = OpenAIResumeParser._clean_string(value)
        if not text:
            return None
        digits = re.sub(r"[^\d+]", "", text)
        digit_count = len(re.sub(r"\D", "", digits))
        return digits if digit_count >= 7 else None

    @staticmethod
    def _normalize_social_url(value: Any, allowed_hosts: tuple[str, ...]) -> Optional[str]:
        text = OpenAIResumeParser._clean_string(value)
        if not text:
            return None
        cleaned = text.strip().strip('"').strip("'")
        if cleaned.startswith("www."):
            cleaned = "https://" + cleaned
        if not cleaned.startswith(("http://", "https://")) and ("linkedin.com" in cleaned or "github.com" in cleaned):
            cleaned = "https://" + cleaned.lstrip("/")
        if any(host in cleaned.lower() for host in allowed_hosts):
            return cleaned
        return None

    @staticmethod
    def _normalize_skills(value: Any) -> List[str]:
        if value is None:
            return []
        if isinstance(value, str):
            raw_skills = [s.strip() for s in re.split(r",|\n|;", value) if s.strip()]
        elif isinstance(value, list):
            raw_skills = [str(s).strip() for s in value if str(s).strip()]
        else:
            return []
        normalized: List[str] = []
        seen = set()
        for skill in raw_skills:
            key = skill.lower().strip()
            key = SKILL_NORMALIZATION.get(key, key)
            key = re.sub(r"\s+", " ", key).strip()
            if key and key not in seen:
                seen.add(key)
                normalized.append(key)
        return normalized

    @staticmethod
    def _normalize_experience(value: Any) -> Optional[float]:
        if value is None:
            return None
        if isinstance(value, (int, float)):
            return float(value)
        text = OpenAIResumeParser._clean_string(value)
        if not text:
            return None
        match = re.search(r"\d+(?:\.\d+)?", text)
        return float(match.group(0)) if match else None

    @staticmethod
    def _normalize_projects(value: Any) -> List[Dict[str, Any]]:
        if not value or not isinstance(value, list):
            return []
        projects: List[Dict[str, Any]] = []
        for item in value:
            if not isinstance(item, dict):
                continue
            title = OpenAIResumeParser._clean_string(item.get("title"))
            description = OpenAIResumeParser._clean_string(item.get("description"))
            technologies = OpenAIResumeParser._normalize_skills(item.get("technologies"))
            if title or description or technologies:
                projects.append({"title": title, "description": description, "technologies": technologies})
        return projects


def discover_resume_files(input_dir: Path, recursive: bool = True) -> List[Path]:
    files = [p for p in (input_dir.rglob("*") if recursive else input_dir.glob("*")) if p.is_file()]
    known = [f for f in files if f.suffix.lower() in SUPPORTED_EXTENSIONS]
    unknown = [f for f in files if f.suffix.lower() not in SUPPORTED_EXTENSIONS]
    return sorted(known) + sorted(unknown)


def parse_single_file(file_path: Path, local_parser: RuleBasedResumeParser, llm_parser: Optional[Any] = None) -> Dict[str, Any]:
    try:
        text = ResumeTextExtractor.extract_text(file_path)
        if not text.strip():
            raise ValueError("No extractable text found")
        local_data = local_parser.parse_resume(text)
        result: Dict[str, Any] = {
            "source_file": str(file_path),
            "status": "success",
            "data": local_data,
            "error": None,
            "parse_method": "local",
        }
        if llm_parser is not None and should_use_llm_fallback(local_data, text):
            try:
                llm_data = llm_parser.parse_resume(text)
                result["data"] = merge_resume_records(local_data, llm_data)
                result["parse_method"] = "local+llm"
            except Exception as exc:
                logging.warning("LLM fallback failed for %s: %s", file_path, exc)
        return result
    except Exception as exc:
        return {
            "source_file": str(file_path),
            "status": "failed",
            "data": None,
            "error": str(exc),
            "parse_method": "local",
        }


def should_use_llm_fallback(parsed_data: Dict[str, Any], raw_text: str) -> bool:
    score = 0
    if parsed_data.get("full_name"):
        score += 2
    if parsed_data.get("email"):
        score += 1
    if parsed_data.get("phone"):
        score += 1
    if parsed_data.get("skills"):
        score += 2
    if parsed_data.get("projects"):
        score += 2
    if parsed_data.get("education"):
        score += 1
    if parsed_data.get("current_role"):
        score += 1
    if parsed_data.get("linkedin_url"):
        score += 1
    if parsed_data.get("github_url"):
        score += 1

    project_signal = bool(re.search(r"\b(projects?|academic projects?|personal projects?|key projects?)\b", raw_text, re.IGNORECASE))
    if project_signal and not parsed_data.get("projects"):
        return True
    return score < 6


def merge_resume_records(local_data: Dict[str, Any], llm_data: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(local_data)
    for key in (
        "full_name",
        "email",
        "phone",
        "linkedin_url",
        "github_url",
        "skills",
        "total_experience_years",
        "education",
        "current_role",
    ):
        value = llm_data.get(key)
        if value not in (None, [], ""):
            merged[key] = value
    merged["projects"] = _merge_projects(local_data.get("projects", []), llm_data.get("projects", []))
    return merged


def _merge_projects(local_projects: List[Dict[str, Any]], llm_projects: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    merged: List[Dict[str, Any]] = []
    seen = set()
    for project in list(local_projects or []) + list(llm_projects or []):
        if not isinstance(project, dict):
            continue
        title = _clean_string(project.get("title")) or ""
        description = _clean_string(project.get("description")) or ""
        technologies = tuple(project.get("technologies") or [])
        fingerprint = (title.lower(), description[:120].lower(), technologies)
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        merged.append({"title": title or None, "description": description or None, "technologies": list(technologies)})
    return merged


def _clean_string(value: Any) -> Optional[str]:
    if value is None:
        return None
    if not isinstance(value, str):
        value = str(value)
    value = value.strip()
    return value or None


def parse_resumes_flow(input_dir: Path, output_path: Path) -> int:
    if not input_dir.exists() or not input_dir.is_dir():
        logging.error("Resume input folder does not exist: %s", input_dir)
        return 1

    local_parser = RuleBasedResumeParser()
    llm_parser = None
    api_key = os.getenv("LLM_API_KEY")
    api_base = os.getenv("LLM_API_BASE")
    llm_model = os.getenv("LLM_MODEL", "gpt-4o-mini")

    if api_key and OpenAI is not None:
        try:
            llm_parser = OpenAIResumeParser(
                LLMConfig(api_key=api_key, model=llm_model, api_base=api_base)
            )
            logging.info("Resume parser: Hybrid mode enabled with model %s", llm_model)
        except Exception as exc:
            logging.warning("Resume parser: OpenAI fallback disabled: %s", exc)
    else:
        logging.info("Resume parser: Running local parser only")

    files = discover_resume_files(input_dir, recursive=True)
    if not files:
        logging.warning("No files found in resume directory: %s", input_dir)
        return 1

    logging.info("Resume parser: Found %d files. Starting parsing...", len(files))
    results = [parse_single_file(path, local_parser, llm_parser) for path in files]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "total_files": len(results),
        "successful": sum(1 for r in results if r["status"] == "success"),
        "failed": sum(1 for r in results if r["status"] == "failed"),
        "results": results,
    }
    output_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    logging.info("Resume parser: Output written to %s", output_path)
    return 0


def parse_jds_flow(input_dir: Path, output_path: Path) -> int:
    if process_jd_directory is None:
        logging.error("JD parser module is not available")
        return 1

    if not input_dir.exists() or not input_dir.is_dir():
        logging.error("JD input folder does not exist: %s", input_dir)
        return 1

    llm_config = None
    api_key = os.getenv("LLM_API_KEY")
    api_base = os.getenv("LLM_API_BASE")
    llm_model = os.getenv("LLM_MODEL", "gpt-4o-mini")
    jd_llm_mode = os.getenv("JD_LLM_MODE", "primary").strip().lower()
    jd_llm_workers = int(os.getenv("JD_LLM_WORKERS", "4"))

    if api_key and OpenAIJDParser is not None:
        try:
            llm_config = JDLLMConfig(api_key=api_key, model=llm_model, api_base=api_base)
            logging.info(
                "JD parser: LLM mode enabled with model %s (mode=%s, workers=%d)",
                llm_model,
                jd_llm_mode,
                max(1, jd_llm_workers),
            )
        except Exception as exc:
            logging.warning("JD parser: OpenAI fallback disabled: %s", exc)
    else:
        logging.info("JD parser: Running local parser only")

    results = process_jd_directory(
        input_dir,
        recursive=True,
        use_llm=llm_config is not None,
        llm_config=llm_config,
        llm_mode=jd_llm_mode,
        max_workers=max(1, jd_llm_workers),
    )

    if not results:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        logging.error(
            "No JD files found in %s. Add at least one JD file with a supported extension: %s",
            input_dir,
            supported,
        )
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "total_files": len(results),
        "parse_type": "job_descriptions",
        "results": results,
    }
    output_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    logging.info("JD parser: Output written to %s", output_path)
    return 0


def run_matching_flow(
    resume_file: Path,
    jd_file: Path,
    output_file: Path,
    top_k: int,
    no_reranking: bool,
    scoring_mode: str = "hybrid",
    client_context_bundle: Optional[Dict[str, Any]] = None,
    parsed_client_output_file: Optional[Path] = None,
) -> int:
    if VectorPipeline is None:
        logging.error("Matching pipeline is not available. Ensure vector modules are installed.")
        return 1

    if not resume_file.exists():
        logging.error("Resume file not found: %s", resume_file)
        return 1
    if not jd_file.exists():
        logging.error("JD file not found: %s", jd_file)
        return 1

    pinecone_api_key = os.getenv("PINECONE_API_KEY")
    if not pinecone_api_key or pinecone_api_key == "pcsk_xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx":
        logging.warning("PINECONE_API_KEY is missing or placeholder; proceeding with local vector scoring")
        pinecone_api_key = ""

    cloud = os.getenv("PINECONE_CLOUD", "aws")
    region = os.getenv("PINECONE_REGION", "us-east-1")
    use_reranking = os.getenv("USE_PINECONE_RERANKING", "true").lower() == "true"
    use_reranking = use_reranking and not no_reranking

    client_context = client_context_bundle.get("merged_context") if isinstance(client_context_bundle, dict) else None

    pipeline = VectorPipeline(
        pinecone_api_key=pinecone_api_key,
        cloud=cloud,
        region=region,
        use_reranking=use_reranking,
        scoring_mode=scoring_mode,
        client_context=client_context,
    )

    output_file.parent.mkdir(parents=True, exist_ok=True)
    success = pipeline.run_complete_pipeline(
        resume_file=str(resume_file),
        jd_file=str(jd_file),
        output_file=str(output_file),
        top_k=top_k,
    )
    if not success:
        logging.error("Matching pipeline failed")
        return 1

    if client_context_bundle:
        try:
            parsed_client_output = parsed_client_output_file or DEFAULT_PARSED_CLIENT_DATA_FILE
            parsed_client_output.parent.mkdir(parents=True, exist_ok=True)
            parsed_client_output.write_text(
                json.dumps(client_context_bundle, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
        except Exception as exc:
            logging.warning("Matching completed but parsed client data could not be written: %s", exc)

    logging.info("Matching pipeline completed. Output written to %s", output_file)
    return 0


def run_all_flow(args: argparse.Namespace) -> int:
    resume_status = parse_resumes_flow(args.input_dir, args.resume_output)
    if resume_status != 0:
        return resume_status

    jd_status = parse_jds_flow(args.jd_input_dir, args.jd_output)
    if jd_status != 0:
        return jd_status

    # Skip matching if VectorPipeline is not available
    if VectorPipeline is None:
        logging.warning("Matching pipeline is not available. Skipping matching flow. Install vector dependencies to enable matching.")
        return 0

    client_context_bundle = load_client_context(args.client_data_dir)

    return run_matching_flow(
        resume_file=args.resume_output,
        jd_file=args.jd_output,
        output_file=args.output,
        top_k=args.top_k,
        no_reranking=args.no_reranking,
        scoring_mode=args.scoring_mode,
        client_context_bundle=client_context_bundle,
        parsed_client_output_file=args.parsed_client_output,
    )


def main() -> int:
    load_env_file(Path(__file__).with_name(".env"))
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

    parser = argparse.ArgumentParser(
        description="AI Sourcing Agent command router",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "command",
        nargs="?",
        default="all",
        choices=["all", "parse-resumes", "parse-jds", "match"],
        help="Workflow command to run",
    )
    parser.add_argument("--parse-jds", action="store_true", help=argparse.SUPPRESS)

    parser.add_argument("--input-dir", type=Path, default=DEFAULT_INPUT_DIR, help="Resume input directory")
    parser.add_argument(
        "--jd-input-dir",
        type=Path,
        default=Path(__file__).resolve().parent / "jd",
        help="Job description input directory",
    )
    parser.add_argument("--resume-output", type=Path, default=DEFAULT_OUTPUT_FILE, help="Parsed resumes output JSON")
    parser.add_argument(
        "--jd-output",
        type=Path,
        default=Path(__file__).resolve().parent / "results" / "parsed_jds.json",
        help="Parsed JDs output JSON",
    )
    parser.add_argument("--resumes", type=Path, default=DEFAULT_OUTPUT_FILE, help="Parsed resumes JSON path")
    parser.add_argument(
        "--jds",
        type=Path,
        default=Path(__file__).resolve().parent / "results" / "parsed_jds.json",
        help="Parsed JDs JSON path",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).resolve().parent / "results" / "matching_results.json",
        help="Matching results output JSON",
    )
    parser.add_argument(
        "--client-data-dir",
        type=Path,
        default=DEFAULT_CLIENT_DATA_DIR,
        help="Flat client data directory with JSON files",
    )
    parser.add_argument(
        "--parsed-client-output",
        type=Path,
        default=DEFAULT_PARSED_CLIENT_DATA_FILE,
        help="Output JSON file for parsed client data",
    )
    parser.add_argument("--top-k", type=int, default=10, help="Top matches per JD")
    parser.add_argument("--no-reranking", action="store_true", help="Disable Pinecone reranking")
    parser.add_argument(
        "--scoring-mode",
        type=str,
        default="hybrid",
        choices=["vector", "hybrid", "weighted"],
        help="Scoring mode for matching (vector, hybrid, weighted)",
    )

    args = parser.parse_args()

    command = args.command
    if args.parse_jds and command == "all":
        command = "parse-jds"

    if command == "parse-resumes":
        return parse_resumes_flow(args.input_dir, args.resume_output)
    if command == "parse-jds":
        return parse_jds_flow(args.jd_input_dir, args.jd_output)
    if command == "match":
        client_context_bundle = load_client_context(args.client_data_dir)
        return run_matching_flow(
            args.resumes,
            args.jds,
            args.output,
            args.top_k,
            args.no_reranking,
            args.scoring_mode,
            client_context_bundle,
            args.parsed_client_output,
        )
    return run_all_flow(args)


if __name__ == "__main__":
    sys.exit(main())
